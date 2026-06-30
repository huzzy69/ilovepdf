import os
import tempfile
import docx
from docx import Document
from rapidocr_onnxruntime import RapidOCR

def convert_image_to_word(image_path: str) -> str:
    """
    Performs OCR on an image and converts the recognized text into a structured Word document (.docx).
    """
    engine = RapidOCR()
    result, elapse = engine(image_path)
    
    doc = Document()
    
    # Configure margins for a clean look
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    if not result:
        doc.add_paragraph("[No text recognized in the image]")
        fd, docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(docx_path)
        return docx_path

    # Sort results top-to-bottom based on y-coordinate of top-left corner
    result_sorted = sorted(result, key=lambda x: x[0][0][1])
    
    # Group lines into paragraphs
    paragraphs = []
    current_paragraph_lines = []
    last_y_bottom = None
    last_x_left = None
    
    # Estimate typical line height
    heights = []
    for line in result_sorted:
        box, text, conf = line
        y_coords = [p[1] for p in box]
        heights.append(max(y_coords) - min(y_coords))
    
    avg_line_height = sum(heights) / len(heights) if heights else 15
    
    for line in result_sorted:
        box, text, conf = line
        text = text.strip()
        if not text:
            continue
            
        x_coords = [p[0] for p in box]
        y_coords = [p[1] for p in box]
        
        y_top = min(y_coords)
        y_bottom = max(y_coords)
        x_left = min(x_coords)
        x_right = max(x_coords)
        
        if not current_paragraph_lines:
            current_paragraph_lines.append(text)
        else:
            vertical_gap = y_top - last_y_bottom
            is_new_para = False
            
            # If vertical gap is large, it's a new paragraph
            if vertical_gap > 1.5 * avg_line_height:
                is_new_para = True
            # If the vertical gap is significantly negative, it might be a column transition or OCR layout issue
            elif vertical_gap < -0.5 * avg_line_height:
                is_new_para = True
            # If horizontal alignment shifts significantly
            elif abs(x_left - last_x_left) > 3.0 * avg_line_height:
                is_new_para = True
            
            # Common markers for list items or headers
            # E.g. "Q1.", "Q2.", "1.", "2.", "•", "-"
            if text.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q6", "Q7", "Q8", "Q9", "Question")):
                is_new_para = True
            elif text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.")):
                is_new_para = True
            elif text.startswith(("-", "*", "•")):
                is_new_para = True
                
            if is_new_para:
                paragraphs.append(" ".join(current_paragraph_lines))
                current_paragraph_lines = [text]
            else:
                current_paragraph_lines.append(text)
                
        last_y_bottom = y_bottom
        last_x_left = x_left
        
    if current_paragraph_lines:
        paragraphs.append(" ".join(current_paragraph_lines))
        
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
            
        # Format headings or headers
        is_heading = False
        if para_text.startswith(("Q1", "Q2", "Q3", "Q4", "Q5", "Q6", "Q7", "Q8", "Q9", "Question")):
            is_heading = True
        elif len(para_text) < 60 and ("truth table" in para_text.lower() or para_text.lower().startswith(("table", "note", "diagram"))):
            is_heading = True
            
        if is_heading:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.bold = True
            run.font.size = docx.shared.Pt(12)
        else:
            doc.add_paragraph(para_text)
            
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(docx_path)
    return docx_path
